from io import BytesIO
from pathlib import Path


FORMATOS_SOPORTADOS = ("txt", "md", "csv", "pdf", "docx")


def decodificar_texto(contenido: bytes) -> str:
    """Convierte archivos de texto a string probando codificaciones comunes."""
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return contenido.decode(encoding)
        except UnicodeDecodeError:
            continue

    return contenido.decode("utf-8", errors="ignore")


def extraer_pdf(contenido: bytes) -> str:
    """Extrae texto de un archivo PDF usando pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("falta instalar pypdf para leer pdf") from exc

    lector = PdfReader(BytesIO(contenido))
    paginas = [pagina.extract_text() or "" for pagina in lector.pages]
    texto = "\n".join(paginas).strip()
    if not texto:
        raise ValueError("no se pudo extraer texto del pdf")
    return texto


def extraer_docx(contenido: bytes) -> str:
    """Extrae texto de un archivo DOCX usando python-docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("falta instalar python-docx para leer docx") from exc

    documento = Document(BytesIO(contenido))
    partes = [parrafo.text for parrafo in documento.paragraphs if parrafo.text.strip()]

    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [celda.text.strip() for celda in fila.cells if celda.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))

    texto = "\n".join(partes).strip()
    if not texto:
        raise ValueError("no se pudo extraer texto del docx")
    return texto


def extraer_texto_archivo(nombre_archivo: str, contenido: bytes) -> str:
    """Extrae texto segun la extension del archivo subido."""
    extension = Path(nombre_archivo).suffix.lower().lstrip(".")

    if extension in {"txt", "md", "csv"}:
        texto = decodificar_texto(contenido).strip()
    elif extension == "pdf":
        texto = extraer_pdf(contenido)
    elif extension == "docx":
        texto = extraer_docx(contenido)
    else:
        formatos = ", ".join(FORMATOS_SOPORTADOS)
        raise ValueError(f"formato no soportado, usa {formatos}")

    if not texto:
        raise ValueError("el archivo no tiene texto util")

    return texto
